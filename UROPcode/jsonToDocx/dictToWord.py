from docx import Document

import json

with open('ACS_photostability_dict.json', 'r') as file:
    data = file.read() 
     

dic  = json.loads(data)

document = Document()

count = 0
for key,value in dic.items():
    if value != 0:
        document.add_heading(key)
        document.add_paragraph(value)
    else:
        count += 1
    
document.save('ACS_photostability.docx')
print(len(dic)-count)